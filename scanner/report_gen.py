import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fpdf import FPDF, HTMLMixin
from typing import Dict, Optional, Union
import logging
from .ai_helper import AIHelper
import matplotlib.pyplot as plt
from collections import defaultdict


def to_latin1(text):
    if not isinstance(text, str):
        text = str(text)
    return text.encode("latin-1", "replace").decode("latin-1")

class HTMLPDF(FPDF, HTMLMixin):
    pass

class ReportGenerator:
    def __init__(self, output_dir: str = "reports", use_ai: bool = False, api_token: str = None):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        self.ai_helper = AIHelper(api_token=api_token) if use_ai and api_token else None
        self.logger = logging.getLogger(__name__)
        self.severity_mapping = {
            'critical': 'Critical',
            'high': 'High',
            'medium': 'Medium',
            'low': 'Low'
        }

    def _normalize_severity(self, severity: str) -> str:
        """Normalize severity to standard title case format"""
        if not severity:
            return "Unknown"
        severity = str(severity).strip().upper()
        
        if severity.startswith('CRIT'):
            return "Critical"
        elif severity.startswith('HIGH'):
            return "High"
        elif severity.startswith('MED'):
            return "Medium"
        elif severity.startswith('LOW'):
            return "Low"
        return severity.title()

    def _normalize_vulnerability_name(self, vuln: Dict) -> str:
        """Normalize vulnerability names consistently across all formats"""
        vuln_name = vuln.get('name', 'Unnamed Vulnerability')
        current_desc = str(vuln.get('description', '')).lower()
        
        if any(keyword in current_desc 
              for keyword in ['outdated', 'old version', 'deprecated']):
            if not vuln_name.startswith(('Outdated ', 'Deprecated ')):
                vuln_name = f"Outdated {vuln_name}"
        
        elif 'CVE' in vuln_name.upper():
            if not vuln_name.startswith('CVE-'):
                cve_id = next((s for s in vuln_name.split() if 'CVE-' in s.upper()), None)
                if cve_id:
                    vuln_name = f"{cve_id.upper()} - {vuln_name.replace(cve_id, '').strip()}"
                else:
                    vuln_name = f"CVE: {vuln_name}"
        
        if 'bootstrap' in vuln_name.lower() and 'CVE' not in vuln_name.upper():
            if 'outdated' not in vuln_name.lower():
                vuln_name = f"Outdated {vuln_name}"
        
        return vuln_name

    def _get_cvss_score(self, vuln: Dict) -> float:
        """Enhanced CVSS score extraction with better fallback logic"""
        # Try explicit cvss_score first
        cvss = vuln.get('cvss_score')
        if cvss not in [None, 'N/A', '']:
            try:
                return float(cvss)
            except (ValueError, TypeError):
                pass
        
        # Try to extract from description if contains CVSS
        desc = str(vuln.get('description', ''))
        if 'CVSS' in desc:
            try:
                parts = desc.split('CVSS')
                score_part = parts[1].split()[0]
                return float(score_part.strip('.:'))
            except (IndexError, ValueError):
                pass
        
        # Fallback to severity mapping
        severity = str(vuln.get('severity', '')).lower()
        if 'critical' in severity:
            return 9.0
        elif 'high' in severity:
            return 7.5
        elif 'medium' in severity:
            return 5.0
        elif 'low' in severity:
            return 2.5
        
        return 0.0

    def _format_cvss_display(self, cvss_score: Union[str, float]) -> str:
        """Format CVSS score with severity rating"""
        if cvss_score in [None, 'N/A', '']:
            return "CVSS: N/A"
        
        try:
            score = float(cvss_score)
            if score >= 9.0:
                severity = "Critical"
            elif score >= 7.0:
                severity = "High"
            elif score >= 4.0:
                severity = "Medium"
            else:
                severity = "Low"
            return f"CVSS: {score:.1f} ({severity})"
        except ValueError:
            return f"CVSS: {cvss_score}"
        
    def _group_by_severity(self, findings):
        """Return dict: {'Critical': [...], 'High': [...], 'Medium': [...], 'Low': [...], 'Other': [...]}"""
        buckets = {"Critical": [], "High": [], "Medium": [], "Low": [], "Other": []}
        for f in findings or []:
            sev = self._normalize_severity(f.get("severity", "Unknown"))
            if sev not in buckets:
                sev = "Other"
            buckets[sev].append(f)
        return buckets

    def _sla(self, severity: str) -> str:
        """SLA guidance by severity (used in Word and PDF)."""
        sev = (severity or "").lower()
        if sev.startswith("crit") or sev == "critical":
            return "Fix within 7 days"
        if sev.startswith("high"):
            return "Fix within 7 days"
        if sev.startswith("med"):
            return "Fix within 30 days"
        if sev.startswith("low"):
            return "Fix within 90 days"
        return "Fix as part of routine backlog"

    def _generate_charts(self, findings, output_dir=None):
        """
        Create pie (severity) and bar (category) charts.
        Returns (pie_path, bar_path) or (None, None) on failure.
        """
        try:
            outdir = output_dir or self.output_dir or "reports"
            os.makedirs(outdir, exist_ok=True)

            # --- Pie: by severity
            severities = ["Critical", "High", "Medium", "Low", "Info"]
            severity_colors = {
                "Critical": "#d62728",  # red
                "High": "#ff7f0e",      # orange
                "Medium": "#ffbf00",    # yellow
                "Low": "#1f77b4",       # blue
                "Info": "#2ca02c",      # green
            }
            counts = {s: 0 for s in severities}

            for f in findings or []:
                sev = self._normalize_severity(f.get("severity", "Unknown"))
                if sev in counts:
                    counts[sev] += 1

            labels = [s for s, c in counts.items() if c > 0]
            sizes = [counts[s] for s in labels]
            colors = [severity_colors[s] for s in labels]

            pie_path = None
            if sizes:
                pie_path = os.path.join(outdir, "severity_pie.png")
                plt.figure()
                wedges, texts, autotexts = plt.pie(
                    sizes,
                    colors=colors,
                    autopct="%1.1f%%",
                    startangle=140,
                    textprops={"color": "w"}
                )
                plt.legend(
                    wedges,
                    labels,
                    title="Severity",
                    loc="center left",
                    bbox_to_anchor=(1, 0, 0.5, 1)
                )
                plt.title("Vulnerabilities by Severity")
                plt.tight_layout()
                plt.savefig(pie_path, bbox_inches="tight")
                plt.close()

            # --- Bar: by category
            counts = defaultdict(int)
            for f in findings or []:
                cat = f.get("category") or "Uncategorized"
                counts[cat] += 1

            bar_path = None
            if counts:
                bar_path = os.path.join(outdir, "category_bar.png")
                plt.figure()
                plt.bar(list(counts.keys()), list(counts.values()))
                plt.xticks(rotation=45, ha="right")
                plt.title("Vulnerabilities by Category")
                plt.tight_layout()
                plt.savefig(bar_path, bbox_inches="tight")
                plt.close()

            return pie_path, bar_path
        except Exception as e:
            self.logger.warning(f"Chart generation failed (non-fatal): {e}")
            return None, None

        

    def generate(self, data: Dict, filename: Optional[str] = None, 
            fmt: str = "word", console: bool = False) -> Union[str, Dict]:
        """Generate a report with AI enhancement"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"webscanr_report_{timestamp}"

        # Batch AI enhancement for vulnerabilities
        if self.ai_helper and data.get("vulnerabilities"):
            try:
                from tqdm import tqdm
            except ImportError:
                tqdm = None
            vulns = data["vulnerabilities"]
            # Prepare vulnerabilities for batch prompt
            prepped_vulns = []
            for vuln in vulns:
                v = vuln.copy()
                v['name'] = self._normalize_vulnerability_name(v)
                v['severity'] = self._normalize_severity(v.get("severity"))
                v['cvss_score'] = self._get_cvss_score(v)
                prepped_vulns.append(v)

            # Build batch prompt for AI
            prompt = """
You are a cybersecurity expert. For each vulnerability below, provide:
1. A clear, concise description (max 5 sentences).
2. Actionable remediation steps (as a bullet list).
3. A business impact summary (max 3 sentences).
Format your response as a JSON array, one object per vulnerability, with keys: description, remediation, business_impact. Keep the order the same as input.
Vulnerabilities:
"""
            for i, v in enumerate(prepped_vulns, 1):
                prompt += f"\n{i}. Name: {v['name']}\nSeverity: {v['severity']}\nCVSS: {v['cvss_score']}\nDescription: {v.get('description', '')}\n"
            prompt += "\n---\nJSON:"

            # Call AI once for all vulnerabilities using Qwen/Qwen3-32B:cerebras
            ai_response = None
            try:
                if tqdm:
                    iterator = tqdm([1], desc="AI-enhancing vulnerabilities (batch)")
                    for _ in iterator:
                        ai_response = self.ai_helper.generate_batch(
                            prompt,
                            model="Qwen/Qwen3-32B:cerebras",
                            base_url="https://router.huggingface.co/v1"
                        )
                else:
                    ai_response = self.ai_helper.generate_batch(
                        prompt,
                        model="Qwen/Qwen3-32B:cerebras",
                        base_url="https://router.huggingface.co/v1"
                    )
            except Exception as e:
                self.logger.error(f"Failed to batch enhance vulnerabilities: {str(e)}")
                ai_response = None

            # Parse AI response and assign to vulnerabilities (robust JSON extraction)
            import json as _json
            import re
            enhanced_vulns = []
            parsed = []
            if ai_response:
                try:
                    # Use greedy regex to extract the largest JSON array from the response
                    match = re.search(r'\[.*\]', ai_response, re.DOTALL)
                    if match:
                        json_str = match.group(0)
                        try:
                            parsed = _json.loads(json_str)
                        except Exception as e:
                            # Try to parse as much as possible if the array is truncated
                            self.logger.warning(f"AI JSON appears truncated or malformed, attempting partial parse: {str(e)}")
                            # Try to parse as many valid top-level objects as possible
                            # Only include objects that start and end with curly braces at the top level
                            items = re.findall(r'\{(?:[^{}]|\{[^{}]*\})*\}', json_str, re.DOTALL)
                            partial = []
                            for item in items:
                                item = item.strip()
                                # Ignore objects that are too short to be valid
                                if not (item.startswith('{') and item.endswith('}')) or len(item) < 10:
                                    continue
                                # Remove trailing commas (common in truncated arrays)
                                if item.endswith(','):
                                    item = item[:-1]
                                try:
                                    partial.append(_json.loads(item))
                                except Exception as obj_err:
                                    self.logger.warning(f"Skipping malformed object in AI JSON: {obj_err}\nObject: {item}")
                                    continue
                            if partial:
                                parsed = partial
                            else:
                                raise
                    else:
                        raise ValueError("No JSON array found in AI response.")
                except Exception as e:
                    self.logger.error(f"Failed to parse AI batch response: {str(e)}\nRaw AI response:\n{ai_response}")
                    parsed = []

            # Use as many valid AI-enhanced objects as possible, fallback to default for the rest
            for i, vuln in enumerate(prepped_vulns):
                enhanced = vuln.copy()
                if parsed and i < len(parsed):
                    ai_vuln = parsed[i]
                    enhanced["description"] = ai_vuln.get("description", vuln.get("description", ""))
                    enhanced["remediation"] = ai_vuln.get("remediation", vuln.get("remediation", []))
                    enhanced["business_impact"] = ai_vuln.get("business_impact", vuln.get("business_impact", ""))
                # If parsed is shorter than prepped_vulns, fallback to default for the rest
                enhanced_vulns.append(enhanced)
            enhanced_data = data.copy()
            enhanced_data["vulnerabilities"] = enhanced_vulns
            data = enhanced_data
        
        filename = os.path.join(self.output_dir, filename)
        
        if fmt == "json":
            return self._generate_json(data, filename + ".json")
        elif fmt == "pdf":
            return self._generate_pdf(data, filename + ".pdf")
        elif fmt == "terminal":
            return self._generate_terminal(data)
        else:
            return self._generate_word(data, filename + ".docx")

    def _generate_terminal(self, data: Dict) -> Dict:
        terminal_data = {
            "metadata": {
                "scan_date": data.get("timestamp"),
                "target_url": data.get("url"),
                "total_vulnerabilities": len(data.get("vulnerabilities", []))
            },
            "vulnerabilities": []
        }

        vulns = data.get("vulnerabilities", [])
        if vulns:
            severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
            vulns_sorted = sorted(
                vulns,
                key=lambda x: severity_order.get(self._normalize_severity(x.get("severity", "Unknown")), 4)
            )

            for vuln in vulns_sorted:
                entry = {
                    "name": self._normalize_vulnerability_name(vuln),
                    "severity": self._normalize_severity(vuln.get("severity")),
                    "cvss_score": self._format_cvss_display(vuln.get("cvss_score")),
                    "description": vuln.get("description", "No description"),
                    "remediation": vuln.get("remediation", []),
                    "business_impact": vuln.get("business_impact", "N/A")
                }
                
                if "outdated" in str(vuln.get("description", "")).lower():
                    entry["type"] = "Outdated Component"
                    parts = vuln["description"].split()
                    entry["current_version"] = parts[1]
                    entry["latest_version"] = parts[-1].rstrip(".")
                elif "CVE" in vuln.get("name", ""):
                    entry["type"] = "NVD Vulnerability"
                    entry["cve_id"] = vuln["name"].split(" - ")[-1]
                    entry["sla"] = self._sla(entry["severity"])
                
                terminal_data["vulnerabilities"].append(entry)
        
        return terminal_data

    def _generate_json(self, data: Dict, filepath: str) -> str:
        """Generate JSON report with normalized data"""
        try:
            normalized_data = data.copy()
            if "vulnerabilities" in normalized_data:
                normalized_vulns = []
                for vuln in normalized_data["vulnerabilities"]:
                    normalized_vuln = vuln.copy()
                    normalized_vuln["name"] = self._normalize_vulnerability_name(vuln)
                    normalized_vuln["severity"] = self._normalize_severity(vuln.get("severity"))
                    normalized_vuln["cvss_score"] = self._format_cvss_display(vuln.get("cvss_score"))
                    normalized_vulns.append(normalized_vuln)
                normalized_data["vulnerabilities"] = normalized_vulns

            with open(filepath, "w", encoding="utf-8") as f:
                json.dump(normalized_data, f, indent=2, default=str)
            return filepath
        except Exception as e:
            raise Exception(f"Failed to generate JSON report: {str(e)}")

    def _generate_word(self, data: Dict, filepath: str) -> str:
        """Generate Microsoft Word report with normalized data"""
        try:
            doc = Document()
            
            # Title and metadata
            title = doc.add_heading("Web Application Vulnerability Report", 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            meta = doc.add_paragraph()
            meta.add_run("Scan Date: ").bold = True
            meta.add_run(data.get("timestamp", datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')))
            meta.add_run("\nTarget URL: ").bold = True
            meta.add_run(data.get("url", "N/A"))
            
            # Executive Summary
            doc.add_heading("Executive Summary", level=1)
            vulns = data.get("vulnerabilities", [])
            summary = doc.add_paragraph()
            
            if not vulns:
                summary.add_run("No vulnerabilities were found during the scan.").italic = True
            else:
                severity_counts = {}
                for vuln in vulns:
                    sev = self._normalize_severity(vuln.get("severity", "Unknown"))
                    severity_counts[sev] = severity_counts.get(sev, 0) + 1
                
                summary.add_run("Scan Results Summary:\n").bold = True
                severity_order = ["Critical", "High", "Medium", "Low"]
                for sev in severity_order:
                    count = severity_counts.pop(sev, 0)
                    if count > 0:
                        summary.add_run(f"- {sev}: {count}\n")
                for sev, count in severity_counts.items():
                    summary.add_run(f"- {sev}: {count}\n")
            
            # --- NEW: Insert charts in Executive Summary (non-fatal if missing) ---
            try:
                pie_path = self._generate_charts(vulns, output_dir=os.path.dirname(filepath))
                if pie_path and os.path.exists(pie_path):
                    doc.add_paragraph()
                    doc.add_picture(pie_path, width=Inches(4.5))
            except Exception as _e:
                # Do not fail the whole report if charts fail
                pass


            # Detailed Findings
            if vulns:
                doc.add_heading("Detailed Findings", level=1)
                severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
                vulns_sorted = sorted(
                    vulns,
                    key=lambda x: severity_order.get(self._normalize_severity(x.get("severity", "Unknown")), 4)
                )
                
                for i, vuln in enumerate(vulns_sorted, 1):
                    vuln_name = self._normalize_vulnerability_name(vuln)
                    doc.add_heading(f"{i}. {vuln_name}", level=2)
                    
                    # Severity and CVSS
                    severity = doc.add_paragraph()
                    severity.add_run("Severity: ").bold = True
                    sev_text = self._normalize_severity(vuln.get("severity", "Unknown"))
                    sev_run = severity.add_run(sev_text)
                    
                    severity.add_run("\nCVSS Score: ").bold = True
                    cvss_text = self._format_cvss_display(vuln.get("cvss_score"))
                    severity.add_run(cvss_text)

                                        # --- NEW: Category (if available) ---
                    category = vuln.get("category")
                    if category:
                        catp = doc.add_paragraph()
                        catp.add_run("Category: ").bold = True
                        catp.add_run(str(category))

                    # --- NEW: SLA guidance ---
                    slap = doc.add_paragraph()
                    slap.add_run("SLA: ").bold = True
                    slap.add_run(self._sla(sev_text))

                    
                    # Color coding
                    sev_text_lower = sev_text.lower()
                    if "critical" in sev_text_lower:
                        sev_run.font.color.rgb = RGBColor(255, 0, 0)
                    elif "high" in sev_text_lower:
                        sev_run.font.color.rgb = RGBColor(220, 50, 0)
                    elif "medium" in sev_text_lower:
                        sev_run.font.color.rgb = RGBColor(255, 165, 0)
                    elif "low" in sev_text_lower:
                        sev_run.font.color.rgb = RGBColor(0, 128, 0)
                    else:
                        sev_run.font.color.rgb = RGBColor(128, 128, 128)
                    
                    # Description
                    doc.add_paragraph("Description:", style="Heading 3")
                    doc.add_paragraph(vuln.get("description", "No description provided."))
                    
                    # Proof of Concept
                    doc.add_paragraph("Proof of Concept:", style="Heading 3")
                    poc = vuln.get("poc", "N/A")
                    if isinstance(poc, list):
                        poc = "\n".join(poc)
                    doc.add_paragraph(poc)
                    
                    # Remediation with references
                    doc.add_paragraph("Remediation:", style="Heading 3")
                    remediation = vuln.get("remediation", [])
                    if remediation:
                        if isinstance(remediation, str):
                            doc.add_paragraph(remediation)
                        else:
                            for step in remediation:
                                # Highlight official references
                                if any(ref in step.lower() for ref in ['https://', 'http://', 'reference:', 'see:']):
                                    p = doc.add_paragraph(style="List Bullet")
                                    runner = p.add_run(step.split('http')[0])
                                    runner.bold = True
                                    if 'http' in step:
                                        url_part = 'http' + step.split('http')[1]
                                        p.add_run(url_part).italic = True
                                else:
                                    doc.add_paragraph(step, style="List Bullet")
                    else:
                        doc.add_paragraph("No specific remediation steps provided.")
                    
                    # Business Impact
                    doc.add_paragraph("Business Impact:", style="Heading 3")
                    impact = vuln.get("business_impact", "N/A")
                    if not impact or impact == "N/A":
                        impact = "No business impact analysis provided."
                    doc.add_paragraph(impact)
                    
                    doc.add_paragraph()
            
            doc.save(filepath)
            return filepath
            
        except Exception as e:
            raise Exception(f"Failed to generate Word report: {str(e)}")

    def _generate_pdf(self, data: Dict, filepath: str) -> str:
        """Generate PDF report with normalized data"""
        try:
            pdf = HTMLPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            
            # Title and metadata
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(0, 10, to_latin1("Web Application Vulnerability Report"), ln=True, align='C')
            pdf.set_font("Arial", '', 12)
            pdf.cell(0, 10, to_latin1(f"Scan Date: {data.get('timestamp', datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC'))}"), ln=True)
            pdf.cell(0, 10, to_latin1(f"Target URL: {data.get('url', 'N/A')}"), ln=True)
            
            # Executive Summary
            pdf.ln(10)
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, to_latin1("Executive Summary"), ln=True)
            
            vulns = data.get("vulnerabilities", [])
            pdf.set_font("Arial", '', 12)
            
            if not vulns:
                pdf.cell(0, 10, to_latin1("No vulnerabilities were found during the scan."), ln=True)
            else:
                severity_counts = {}
                for vuln in vulns:
                    sev = self._normalize_severity(vuln.get("severity", "Unknown"))
                    severity_counts[sev] = severity_counts.get(sev, 0) + 1
                
                pdf.cell(0, 10, to_latin1("Scan Results Summary:"), ln=True)
                severity_order = ["Critical", "High", "Medium", "Low"]
                for sev in severity_order:
                    count = severity_counts.pop(sev, 0)
                    if count > 0:
                        pdf.cell(0, 10, to_latin1(f"- {sev}: {count}"), ln=True)
                for sev, count in severity_counts.items():
                    pdf.cell(0, 10, to_latin1(f"- {sev}: {count}"), ln=True)
            
                        # --- NEW: Insert charts in Executive Summary ---
                try:
                    pie_path, bar_path = self._generate_charts(vulns, output_dir=os.path.dirname(filepath))
                    if pie_path and os.path.exists(pie_path):
                        pdf.ln(3)
                        pdf.image(pie_path, w=120)
                except Exception as _e:
                    # Non-fatal
                    pass


            # Detailed Findings
            if vulns:
                pdf.add_page()
                pdf.set_font("Arial", 'B', 14)
                pdf.cell(0, 10, to_latin1("Detailed Findings"), ln=True)
                
                severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
                vulns_sorted = sorted(
                    vulns,
                    key=lambda x: severity_order.get(self._normalize_severity(x.get("severity", "Unknown")), 4)
                )
                
                for i, vuln in enumerate(vulns_sorted, 1):
                    pdf.ln(5)
                    pdf.set_font("Arial", 'B', 12)
                    
                    # Vulnerability title
                    vuln_name = self._normalize_vulnerability_name(vuln)
                    severity = self._normalize_severity(vuln.get("severity", "Unknown"))
                    
                    # Severity color
                    if "critical" in severity.lower():
                        pdf.set_text_color(255, 0, 0)
                    elif "high" in severity.lower():
                        pdf.set_text_color(220, 50, 0)
                    elif "medium" in severity.lower():
                        pdf.set_text_color(255, 165, 0)
                    elif "low" in severity.lower():
                        pdf.set_text_color(0, 128, 0)
                    else:
                        pdf.set_text_color(128, 128, 128)
                    
                    pdf.multi_cell(0, 10, to_latin1(f"{i}. {vuln_name}"))
                    pdf.set_text_color(0, 0, 0)  # Reset to black
                    
                    # Severity and CVSS
                    pdf.set_font("Arial", 'B', 12)
                    pdf.cell(40, 10, to_latin1("Severity:"))
                    pdf.set_font("Arial", '', 12)
                    pdf.cell(0, 10, to_latin1(severity), ln=True)
                    
                    pdf.set_font("Arial", 'B', 12)
                    pdf.cell(40, 10, to_latin1("CVSS Score:"))
                    pdf.set_font("Arial", '', 12)
                    pdf.cell(0, 10, to_latin1(self._format_cvss_display(vuln.get("cvss_score"))), ln=True)
                    
                    # Description
                    pdf.set_font("Arial", 'B', 12)
                    pdf.cell(0, 10, to_latin1("Description:"), ln=True)
                    pdf.set_font("Arial", '', 12)
                    pdf.multi_cell(0, 10, to_latin1(vuln.get("description", "No description provided.")))
                    
                    # Proof of Concept
                    pdf.set_font("Arial", 'B', 12)
                    pdf.cell(0, 10, to_latin1("Proof of Concept:"), ln=True)
                    pdf.set_font("Arial", '', 12)
                    poc = vuln.get("poc", "N/A")
                    if isinstance(poc, list):
                        poc = "\n".join(poc)
                    pdf.multi_cell(0, 10, to_latin1(poc))
                    
                    # Remediation with references
                    pdf.set_font("Arial", 'B', 12)
                    pdf.cell(0, 10, to_latin1("Remediation Steps:"), ln=True)
                    pdf.set_font("Arial", '', 12)
                    remediation = vuln.get("remediation", [])
                    if remediation:
                        if isinstance(remediation, str):
                            pdf.multi_cell(0, 10, to_latin1(remediation))
                        else:
                            for step in remediation:
                                pdf.cell(10)
                                if any(ref in step.lower() for ref in ['https://', 'http://', 'reference:']):
                                    parts = step.split('http')
                                    pdf.set_font("Arial", 'B', 12)
                                    pdf.multi_cell(0, 10, to_latin1(parts[0]))
                                    pdf.set_font("Arial", 'I', 12)
                                    pdf.multi_cell(0, 10, to_latin1('http' + parts[1] if len(parts) > 1 else ''))
                                    pdf.set_font("Arial", '', 12)
                                else:
                                    pdf.multi_cell(0, 10, to_latin1(f"- {step}"))
                    else:
                        pdf.multi_cell(0, 10, to_latin1("No specific remediation steps provided."))
                    
                    # Business Impact
                    pdf.set_font("Arial", 'B', 12)
                    pdf.cell(0, 10, to_latin1("Business Impact:"), ln=True)
                    pdf.set_font("Arial", '', 12)
                    pdf.multi_cell(0, 10, to_latin1(vuln.get("business_impact", "N/A")))
                    
                    pdf.ln(5)

                    # --- NEW: Category (if any) ---
                    category = vuln.get("category")
                    if category:
                        pdf.set_font("Arial", 'B', 12)
                        pdf.cell(40, 10, to_latin1("Category:"))
                        pdf.set_font("Arial", '', 12)
                        pdf.cell(0, 10, to_latin1(str(category)), ln=True)

                    # --- NEW: OWASP (if any) ---
                    owasp = vuln.get("owasp") or vuln.get("owasp_top10")
                    if owasp:
                        pdf.set_font("Arial", 'B', 12)
                        pdf.cell(40, 10, to_latin1("OWASP:"))
                        pdf.set_font("Arial", '', 12)
                        pdf.cell(0, 10, to_latin1(str(owasp)), ln=True)

                    # --- NEW: SLA guidance ---
                    pdf.set_font("Arial", 'B', 12)
                    pdf.cell(40, 10, to_latin1("SLA:"))
                    pdf.set_font("Arial", '', 12)
                    pdf.cell(0, 10, to_latin1(self._sla(severity)), ln=True)

            
            pdf.output(filepath)
            return filepath
            
        except Exception as e:
            raise Exception(f"Failed to generate PDF report: {str(e)}")

    def print_terminal(self, data: Dict, verbose: bool = False) -> None:
        """Print formatted JSON output to terminal for verification."""
        print("\n" + "="*50)
        print("JSON Terminal Output Verification")
        print("="*50)
        
        if verbose:
            print(json.dumps(self._generate_terminal(data), indent=2))
        else:
            terminal_data = self._generate_terminal(data)
            print(f"\nScan Date: {terminal_data['metadata']['scan_date']}")
            print(f"Target URL: {terminal_data['metadata']['target_url']}")
            print(f"Total Vulnerabilities: {terminal_data['metadata']['total_vulnerabilities']}")
            
            if terminal_data['vulnerabilities']:
                print("\nTop Vulnerabilities:")
                for i, vuln in enumerate(terminal_data['vulnerabilities'][:5], 1):
                    print(f"\n{i}. {vuln['name']}")
                    print(f"   Severity: {vuln['severity']}")
                    print(f"   {vuln['cvss_score']}")
                    if "CVE" in vuln.get('name', ''):
                        print(f"   CVE ID: {vuln.get('cve_id', 'N/A')}")
                        print(f"   Description: {vuln.get('description', 'N/A')[:120]}...")
                    else:
                        print(f"   Description: {vuln.get('description', 'N/A')}")
                    print(f"   Business Impact: {vuln.get('business_impact', 'N/A')[:120]}...")
                    print("   Remediation Steps:")
                    print(f"   SLA: {vuln.get('sla', self._sla(vuln['severity']))}")
                    for step in vuln.get('remediation', [])[:3]:
                        print(f"     - {step[:100]}{'...' if len(step) > 100 else ''}")
            else:
                print("\nNo vulnerabilities found.")

        print("\n" + "="*50 + "\n")